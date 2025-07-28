import streamlit as st
import tempfile
import os
import yaml
import json
import subprocess
from pathlib import Path
from typing import Dict, Tuple, Optional, Any
from datetime import datetime
from collections import Counter
import anthropic

# PDF/DOCX extraction imports
try:
    import PyPDF2
    import pdfplumber
    from docx import Document
except ImportError:
    st.error("Required libraries not installed. Please install: pip install PyPDF2 pdfplumber python-docx")

class CVOptimizer:
    """
    A comprehensive CV optimization class that handles:
    - CV text extraction from PDF/DOCX files
    - Template YAML loading
    - Job description processing
    - Prompt template management
    - AI optimization via Haiku 3.5
    """
    
    def __init__(self, templates_dir: str = "../yaml", prompts_dir: str = "../prompt"):
        """
        Initialize the CV Optimizer
        
        Args:
            templates_dir: Directory containing YAML template files (relative to code/ directory)
            prompts_dir: Directory containing prompt template files (relative to code/ directory)
        """
        # Get the directory where this script is located (code/)
        script_dir = Path(__file__).parent
        
        # Create absolute paths relative to the script directory
        self.templates_dir = script_dir / templates_dir
        self.prompts_dir = script_dir / prompts_dir
        self.output_dir = script_dir / "../output"
        
        # Storage variables
        self.cv_text: str = ""
        self.job_description: str = ""
        self.selected_template: str = ""
        self.template_config: Dict[str, Any] = {}
        self.prompt_template: str = ""
        
        # AI client
        self.client = None
        
        # Ensure directories exist
        self.templates_dir.mkdir(exist_ok=True)
        self.prompts_dir.mkdir(exist_ok=True)
        self.output_dir.mkdir(exist_ok=True)
    
    def extract_cv_text(self, uploaded_file) -> Tuple[bool, str]:
        """
        Extract text from uploaded CV file (PDF or DOCX)
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Tuple of (success: bool, extracted_text: str)
        """
        if not uploaded_file:
            return False, "No file uploaded"
        
        try:
            file_type = uploaded_file.type
            
            if file_type == "application/pdf":
                return self._extract_pdf_text(uploaded_file)
            elif file_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ]:
                return self._extract_docx_text(uploaded_file)
            else:
                return False, f"Unsupported file type: {file_type}"
                
        except Exception as e:
            return False, f"Error extracting text: {str(e)}"
    
    def _extract_pdf_text(self, uploaded_file) -> Tuple[bool, str]:
        """Extract text from PDF file using multiple methods for robustness"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_file_path = temp_file.name
            
            extracted_text = ""
            
            # Method 1: Try pdfplumber (more accurate for complex layouts)
            try:
                with pdfplumber.open(temp_file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            extracted_text += page_text + "\n"
                
                if extracted_text.strip():
                    self.cv_text = extracted_text.strip()
                    return True, extracted_text.strip()
            except Exception:
                pass
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(temp_file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        extracted_text += page.extract_text() + "\n"
                
                if extracted_text.strip():
                    self.cv_text = extracted_text.strip()
                    return True, extracted_text.strip()
            except Exception:
                pass
            
            # Clean up temp file
            os.unlink(temp_file_path)
            
            if not extracted_text.strip():
                return False, "Could not extract text from PDF. File might be image-based or corrupted."
            
            return True, extracted_text.strip()
            
        except Exception as e:
            return False, f"PDF extraction error: {str(e)}"
    
    def _extract_docx_text(self, uploaded_file) -> Tuple[bool, str]:
        """Extract text from DOCX file"""
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                temp_file.write(uploaded_file.getvalue())
                temp_file_path = temp_file.name
            
            # Extract text using python-docx
            doc = Document(temp_file_path)
            extracted_text = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text.append(cell.text.strip())
            
            # Clean up temp file
            os.unlink(temp_file_path)
            
            final_text = "\n".join(extracted_text)
            self.cv_text = final_text
            
            return True, final_text
            
        except Exception as e:
            return False, f"DOCX extraction error: {str(e)}"
    
    def set_job_description(self, jd_text: str) -> bool:
        """
        Store job description text
        
        Args:
            jd_text: Job description text
            
        Returns:
            bool: Success status
        """
        if not jd_text or not jd_text.strip():
            return False
        
        self.job_description = jd_text.strip()
        return True
    
    def load_template_config(self, template_name: str) -> Tuple[bool, str]:
        """
        Load YAML configuration for selected template
        
        Args:
            template_name: Name of template (professional, modern, executive, creative)
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            template_file = self.templates_dir / f"{template_name}.yaml"
            
            if not template_file.exists():
                return False, f"Template file not found: {template_file}"
            
            with open(template_file, 'r', encoding='utf-8') as file:
                self.template_config = yaml.safe_load(file)
            
            self.selected_template = template_name
            return True, f"Template '{template_name}' loaded successfully"
            
        except Exception as e:
            return False, f"Error loading template: {str(e)}"
    
    def load_prompt_template(self, prompt_file: str = "prompt_1.txt") -> Tuple[bool, str]:
        """
        Load prompt template from file
        
        Args:
            prompt_file: Name of prompt template file (default: prompt_1.txt)
            
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            prompt_path = self.prompts_dir / prompt_file
            
            if not prompt_path.exists():
                return False, f"Prompt file not found: {prompt_path}"
            
            with open(prompt_path, 'r', encoding='utf-8') as file:
                self.prompt_template = file.read()
            
            return True, "Prompt template loaded successfully"
            
        except Exception as e:
            return False, f"Error loading prompt template: {str(e)}"
    def load_extracted_keywords(self, json_file_path: str) -> Tuple[bool, str]:
        """
        Load extracted ATS keywords from a JSON file
        
        Args:
            json_file_path: Path to the extracted keywords JSON file
        
        Returns:
            Tuple of (success: bool, message: str)
        """
        try:
            with open(json_file_path, 'r', encoding='utf-8') as file:
                self.extracted_keywords_json = file.read()
            return True, "Extracted keywords JSON loaded successfully"
        except Exception as e:
            return False, f"Error loading extracted keywords JSON: {str(e)}"



    def substitute_variables(self) -> Tuple[bool, str]:
        """
        Substitute variables in prompt template
        
        Returns:
            Tuple of (success: bool, final_prompt: str)
        """
        try:
            if not self.prompt_template:
                return False, "No prompt template loaded"
            
            if not self.cv_text:
                return False, "No CV text available"
            
            if not self.job_description:
                return False, "No job description available"
            
            if not self.template_config:
                return False, "No template configuration loaded"
            
            # Prepare substitution variables
            variables = {
                'INSERT_TEMPLATE': json.dumps(self.template_config, indent=2),
                'INSERT_JOB_DESCRIPTION': self.job_description,
                'INSERT_CV': self.cv_text,
                'INSERT_JSON_STRING_HERE': self.extracted_keywords_json  # <-- Add this line
            }
            
            # Substitute variables in prompt
            final_prompt = self.prompt_template.format(**variables)

            

            
            return True, final_prompt
            
        except KeyError as e:
            return False, f"Missing variable in prompt template: {str(e)}"
        except Exception as e:
            return False, f"Error substituting variables: {str(e)}"
    
    def setup_ai_client(self, api_key: str) -> bool:
        """
        Setup Anthropic API client for Haiku 3.5
        
        Args:
            api_key: Anthropic API key
            
        Returns:
            bool: Success status
        """
        try:
            self.client = anthropic.Anthropic(api_key=api_key)
            return True
        except Exception as e:
            st.error(f"Error setting up AI client: {str(e)}")
            return False
    
    def optimize_cv_with_ai(self, prompt: str, max_tokens: int = 4000) -> Tuple[bool, str]:
        try:
            if not self.client:
                return False, "AI client not initialized"

            system_message = (
                    "You are an expert UK recruitment specialist with comprehensive knowledge of UK hiring practices "
                    "across all industries and experience levels. Optimize the CV for the UK job market while maintaining "
                    "complete accuracy, ATS compatibility, UK legal compliance, and RenderCV YAML validity. "
                    "Follow the instructions in  EXACTLY. Output only valid YAML starting with 'cv:' and no explanations."
                )            
                
            response = self.client.messages.create(
                model="claude-3-5-haiku-20241022",
                max_tokens=max_tokens,
                temperature=0.1,
                system=system_message, 
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )

            optimized_cv = response.content[0].text
            return True, optimized_cv

        except Exception as e:
            return False, f"AI optimization error: {str(e)}"

    
    def save_optimized_cv(self, optimized_cv: str, filename: str = None) -> Tuple[bool, str]:
        """
        Save optimized CV to output directory as YAML file
        
        Args:
            optimized_cv: The optimized CV content
            filename: Optional filename, will generate if not provided
            
        Returns:
            Tuple of (success: bool, file_path: str)
        """
        try:
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"optimized_cv_{self.selected_template}_{timestamp}.yaml"
            
            # Ensure filename has .yaml extension
            if not filename.endswith('.yaml'):
                filename = filename.rsplit('.', 1)[0] + '.yaml'
            
            output_path = self.output_dir / filename
            
            # Structure the CV content for YAML format
            cv_data = optimized_cv
            
            with open(output_path, 'w', encoding='utf-8') as file:
                file.write(cv_data)
            
            return True, str(output_path)
            
        except Exception as e:
            return False, f"Error saving CV: {str(e)}"
    
    def generate_pdf_from_yaml(self, yaml_path: str) -> Tuple[bool, str]:
        """
        Generate PDF from YAML file using rendercv CLI tool
        
        Args:
            yaml_path: Path to the YAML file
            
        Returns:
            Tuple of (success: bool, pdf_path_or_error: str)
        """
        try:
            yaml_path = Path(yaml_path)
            
            # Check if YAML file exists
            if not yaml_path.exists():
                return False, f"YAML file not found: {yaml_path}"
            
            # Run the rendercv CLI tool
            result = subprocess.run(
                ["rendercv", "render", str(yaml_path)], 
                check=True, 
                capture_output=True, 
                text=True,
                cwd=yaml_path.parent  # Run in the directory containing the YAML file
            )
            
            output_dir = yaml_path.parent / "rendercv_output"

            # Expected PDF path (rendercv usually creates PDF in same directory)
            # pdf_path = output_dir.with_suffix('.pdf')
            
            # Check if PDF was created
            if output_dir.exists():
                pdf_files = list(output_dir.glob("*.pdf"))
                if not pdf_files:
                    return False, f"No PDF files found in: {output_dir}"

                # If multiple, return the most recently modified one
                pdf_files.sort(key=lambda f: f.stat().st_mtime, reverse=True)
                return True, str(pdf_files[0])

            else:
                # Try alternative naming conventions
                alternative_paths = [
                    yaml_path.parent / f"{yaml_path.stem}.pdf",
                    yaml_path.parent / "output" / f"{yaml_path.stem}.pdf",
                    yaml_path.parent / "rendered" / f"{yaml_path.stem}.pdf"
                ]
                
                for alt_path in alternative_paths:
                    if alt_path.exists():
                        return True, str(alt_path)
                
                return False, f"PDF was not generated. rendercv output: {output_dir} {result.stdout} "
            
        except FileNotFoundError:
            return False, "rendercv CLI tool not found. Please install it with: pip install rendercv"
        except subprocess.CalledProcessError as e:
            return False, f"Failed to generate PDF: {e.stderr if e.stderr else str(e)}"
        except Exception as e:
            return False, f"Unexpected error generating PDF: {str(e)}"
    
    def check_rendercv_availability(self) -> Tuple[bool, str]:
        """
        Check if rendercv CLI tool is available
        
        Returns:
            Tuple of (is_available: bool, message: str)
        """
        try:
            result = subprocess.run(
                ["rendercv", "--version"], 
                check=True, 
                capture_output=True, 
                text=True
            )
            return True, f"rendercv is available: {result.stdout.strip()}"
        except FileNotFoundError:
            return False, "rendercv CLI tool not found. Install with: pip install rendercv"
        except subprocess.CalledProcessError as e:
            return False, f"rendercv found but not working properly: {str(e)}"
        except Exception as e:
            return False, f"Error checking rendercv: {str(e)}"
    
    def get_output_directory(self) -> str:
        """Get the output directory path"""
        return str(self.output_dir)
    
    def list_available_templates(self) -> Dict[str, bool]:
        """
        Check which template files are available
        
        Returns:
            Dict with template names and availability status
        """
        templates = ['professional', 'modern', 'executive', 'creative']
        availability = {}
        
        for template in templates:
            template_file = self.templates_dir / f"{template}.yaml"
            availability[template] = template_file.exists()
        
        return availability
    
    def get_directory_info(self) -> Dict[str, str]:
        """
        Get information about directory structure
        
        Returns:
            Dict with directory paths
        """
        return {
            'templates_dir': str(self.templates_dir),
            'prompts_dir': str(self.prompts_dir),
            'output_dir': str(self.output_dir),
            'templates_exists': self.templates_dir.exists(),
            'prompts_exists': self.prompts_dir.exists(),
            'output_exists': self.output_dir.exists()
        }
        """
        Get current status of all optimization components
        
        Returns:
            Dict with component status
        """
        return {
            'cv_extracted': bool(self.cv_text),
            'job_description_set': bool(self.job_description),
            'template_loaded': bool(self.template_config),
            'prompt_loaded': bool(self.prompt_template),
            'ai_client_ready': bool(self.client)
        }
    
    def reset(self):
        """Reset all stored data"""
        self.cv_text = ""
        self.job_description = ""
        self.selected_template = ""
        self.template_config = {}
        self.prompt_template = ""
    
    def full_optimization_pipeline(self, uploaded_file, jd_text: str, template_name: str, 
                                 api_key: str, prompt_file: str = "prompt_1.txt") -> Tuple[bool, str]:
        """
        Complete optimization pipeline
        
        Args:
            uploaded_file: Streamlit uploaded file
            jd_text: Job description text
            template_name: Template name
            api_key: Anthropic API key
            prompt_file: Prompt template file name
            
        Returns:
            Tuple of (success: bool, result: str)
        """
        try:
            # Step 1: Extract CV text
            success, cv_text = self.extract_cv_text(uploaded_file)
            if not success:
                return False, f"CV extraction failed: {cv_text}"
            
            # Step 2: Set job description
            if not self.set_job_description(jd_text):
                return False, "Invalid job description"
            
            # Step 3: Load template config
            success, msg = self.load_template_config(template_name)
            if not success:
                return False, msg
            
            json_file_path = os.path.join(self.output_dir, "jd_extracted.json")  # or dynamic if needed
            success, msg = self.load_extracted_keywords(json_file_path)
            if not success:
                return False, msg


            # Step 4: Load prompt template
            success, msg = self.load_prompt_template(prompt_file)
            if not success:
                return False, msg
            
            # Step 5: Substitute variables
            success, final_prompt = self.substitute_variables()
            if not success:
                return False, final_prompt
            
            # Step 6: Setup AI client
            if not self.setup_ai_client(api_key):
                return False, "Failed to setup AI client"
            
            # Step 7: Optimize with AI
            success, optimized_cv = self.optimize_cv_with_ai(final_prompt)
            if not success:
                return False, optimized_cv
            
            return True, optimized_cv
            
        except Exception as e:
            return False, f"Pipeline error: {str(e)}"


# Example usage function for Streamlit integration
def integrate_with_streamlit():
    """Example of how to integrate with your Streamlit app"""
    
    # Initialize optimizer
    optimizer = CVOptimizer()
    
    # Example usage in your render functions
    def enhanced_cv_extraction(uploaded_file):
        if uploaded_file:
            success, extracted_text = optimizer.extract_cv_text(uploaded_file)
            
            if success:
                st.success("✅ CV text extracted successfully!")
                st.session_state.cv_text = extracted_text
                st.text_area("📄 Extracted CV Content", value=extracted_text, height=300)
                return extracted_text
            else:
                st.error(f"❌ Failed to extract CV text: {extracted_text}")
                return None
        return None
    
    def enhanced_job_description_processing(jd_text):
        if optimizer.set_job_description(jd_text):
            st.session_state.job_description = jd_text
            return True
        return False
    
    def enhanced_template_loading(template_name):
        success, message = optimizer.load_template_config(template_name)
        if success:
            st.session_state.template_config = optimizer.template_config
            st.success(f"✅ {message}")
        else:
            st.error(f"❌ {message}")
        return success
    
    return optimizer